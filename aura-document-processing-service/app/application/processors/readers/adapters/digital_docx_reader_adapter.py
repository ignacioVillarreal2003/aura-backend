import logging
from pathlib import Path

from docx import Document

from app.application.processors.readers.exceptions.reader_exception import (
    ReaderFileNotFoundException,
    UnsupportedDigitalDOCXFormatException,
    DigitalDOCXReadException,
    DOCXHasNoExtractableTextException
)
from app.application.processors.readers.interfaces.reader_adapter_interface import DocumentReaderInterface

logger = logging.getLogger(__name__)


class DigitalDOCXReaderAdapter(DocumentReaderInterface):
    def can_handle(
            self,
            file_path: Path
    ) -> bool:
        if file_path.suffix.lower() != ".docx":
            return False

        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    return True
            return False

        except Exception:
            return False

    def read(
            self,
            file_path: Path
    ) -> str:
        if not file_path.exists():
            raise ReaderFileNotFoundException(str(file_path))

        logger.info(f"Reading digital DOCX [file={file_path.name}]")

        try:
            doc = Document(file_path)
            # Use \n\n as paragraph separator — consistent with PDF reader and
            # compatible with RecursiveCharacterTextSplitter's default separators
            text_parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            if not text_parts:
                raise DOCXHasNoExtractableTextException()

            logger.info(
                f"Digital DOCX read successfully [file={file_path.name}, paragraphs={len(text_parts)}]"
            )
            return "\n\n".join(text_parts)

        except (ReaderFileNotFoundException, DOCXHasNoExtractableTextException):
            raise
        except Exception as e:
            raise DigitalDOCXReadException(str(file_path), e)
