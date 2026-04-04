import logging
from pathlib import Path
from typing import Optional
from docx import Document

from app.application.processors.readers.exceptions.reader_exception import (
    DigitalDOCXReadException,
    DOCXHasNoExtractableTextException,
    ReaderFileNotFoundException,
    ReaderInitializationException
)
from app.application.processors.readers.interfaces.reader_interface import ReaderInterface
from app.application.processors.readers.reader_settings import ReaderSettings

logger = logging.getLogger(__name__)


class DigitalDOCXReader(ReaderInterface):
    _DOCX_MAGIC = b"PK\x03\x04"

    def __init__(self, reader_settings: Optional[ReaderSettings] = None) -> None:
        self._settings = reader_settings or ReaderSettings()

        try:
            logger.info("DigitalDOCXReader initialized successfully")
        except Exception as e:
            logger.exception("Failed to initialize DigitalDOCXReader")
            raise ReaderInitializationException(
                f"DigitalDOCXReader initialization failed: {e}"
            ) from e

    def can_handle(self, file_path: Path) -> bool:
        if file_path.suffix.lower() != ".docx":
            return False

        if not self._is_valid_docx(file_path):
            return False

        try:
            doc = Document(file_path)
            has_paragraph_text = any(p.text and p.text.strip() for p in doc.paragraphs)
            has_table_text = any(
                cell.text and cell.text.strip()
                for table in doc.tables
                for row in table.rows
                for cell in row.cells
            )
            return has_paragraph_text or has_table_text

        except Exception as e:
            logger.debug("Error during can_handle", extra={"file": file_path.name, "error": str(e)})
            return False

    def read(self, file_path: Path) -> str:
        if not file_path.exists():
            raise ReaderFileNotFoundException(
                "The specified DOCX file does not exist or cannot be accessed."
            )

        logger.info("Reading digital DOCX", extra={"file": file_path.name})

        try:
            doc = Document(file_path)
            text_parts = self._extract_text(doc)

            if not text_parts:
                raise DOCXHasNoExtractableTextException(
                    "The DOCX file does not contain extractable text content."
                )

            logger.info(
                "Digital DOCX read successfully",
                extra={"file": file_path.name, "parts": len(text_parts)}
            )

            return "\n\n".join(text_parts)

        except (ReaderFileNotFoundException, DOCXHasNoExtractableTextException):
            raise
        except Exception as e:
            logger.exception("Error reading digital DOCX", extra={"file": file_path.name})
            raise DigitalDOCXReadException(
                "An unexpected error occurred while reading the digital DOCX file."
            ) from e

    def _extract_text(self, doc: Document) -> list[str]:
        text_parts: list[str] = []

        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text and cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return text_parts

    def _is_valid_docx(self, file_path: Path) -> bool:
        try:
            with open(file_path, "rb") as f:
                return f.read(4).startswith(self._DOCX_MAGIC)
        except Exception as e:
            logger.debug("Failed to validate DOCX magic bytes", extra={"file": file_path.name, "error": str(e)})
            return False
